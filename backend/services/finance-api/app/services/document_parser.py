"""
Document Parser Service
Handles PDF tax returns, bank statements, investment docs, CSV transactions
"""

import io
import re
import csv
import json
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime, date
from dataclasses import dataclass
import PyPDF2
import pdfplumber
import pandas as pd
from docx import Document
import pytesseract
from PIL import Image
import numpy as np

@dataclass
class ParsedTaxReturn:
    tax_year: int
    filing_status: str
    agi: float
    taxable_income: float
    total_tax: float
    withholding: float
    refund_or_owed: float
    
    # Income sources
    wages: float = 0
    interest: float = 0
    dividends: float = 0
    capital_gains: float = 0
    business_income: float = 0
    rental_income: float = 0
    
    # Deductions
    standard_or_itemized: str = "standard"
    total_deductions: float = 0
    mortgage_interest: float = 0
    charitable_donations: float = 0
    state_local_taxes: float = 0
    
    # Credits
    child_tax_credit: float = 0
    education_credits: float = 0
    earned_income_credit: float = 0
    
    # Retirement
    ira_contributions: float = 0
    retirement_distributions: float = 0

@dataclass
class ParsedTransaction:
    date: datetime
    description: str
    amount: float
    category: Optional[str] = None
    account: Optional[str] = None
    transaction_type: str = "debit"  # debit or credit
    tags: List[str] = None
    merchant: Optional[str] = None

@dataclass
class ParsedAccount:
    account_name: str
    account_number: str  # Masked
    account_type: str
    balance: float
    institution: str
    as_of_date: datetime

class DocumentParserService:
    
    # Common tax form field mappings
    TAX_FORM_PATTERNS = {
        '1040': {
            'agi': r'adjusted gross income.*?\$?([\d,]+)',
            'wages': r'wages.*?box 1.*?\$?([\d,]+)',
            'taxable_income': r'taxable income.*?\$?([\d,]+)',
            'total_tax': r'total tax.*?\$?([\d,]+)',
            'withholding': r'federal income tax withheld.*?\$?([\d,]+)',
            'filing_status': r'(single|married filing jointly|married filing separately|head of household)',
        },
        'w2': {
            'wages': r'box 1.*?\$?([\d,]+)',
            'federal_withholding': r'box 2.*?\$?([\d,]+)',
            'social_security_wages': r'box 3.*?\$?([\d,]+)',
            '401k_contributions': r'box 12.*?D.*?\$?([\d,]+)',
        },
        '1099': {
            'interest': r'interest income.*?\$?([\d,]+)',
            'dividends': r'ordinary dividends.*?\$?([\d,]+)',
            'capital_gains': r'capital gain.*?\$?([\d,]+)',
        }
    }
    
    # Transaction categorization rules
    CATEGORY_RULES = {
        'food': ['restaurant', 'grocery', 'uber eats', 'doordash', 'grubhub', 'starbucks', 'mcdonalds'],
        'transportation': ['uber', 'lyft', 'gas', 'shell', 'chevron', 'parking', 'toll'],
        'shopping': ['amazon', 'target', 'walmart', 'best buy', 'home depot'],
        'utilities': ['electric', 'gas', 'water', 'internet', 'comcast', 'at&t', 'verizon'],
        'entertainment': ['netflix', 'spotify', 'hulu', 'movie', 'theater', 'concert'],
        'healthcare': ['pharmacy', 'cvs', 'walgreens', 'doctor', 'hospital', 'dental'],
        'insurance': ['insurance', 'geico', 'state farm', 'allstate', 'progressive'],
        'rent': ['rent', 'lease', 'apartment'],
        'mortgage': ['mortgage', 'home loan', 'property tax', 'hoa'],
    }
    
    async def parse_tax_return(self, file_content: bytes, file_type: str = 'pdf') -> ParsedTaxReturn:
        """
        Parse tax return from PDF or image
        """
        if file_type == 'pdf':
            text = self._extract_pdf_text(file_content)
        else:
            # OCR for scanned documents (hybrid Tesseract + Claude Vision)
            text = await self._extract_image_text(file_content, doc_type='1040')
        
        # Extract tax year - dynamic pattern supports any reasonable tax year
        # Looks for 4-digit years from 2000-2099 (covers all modern tax returns)
        current_year = datetime.now().year
        year_match = re.search(r'\b(20[0-9]{2})\b', text)

        if year_match:
            found_year = int(year_match.group(1))
            # Validate year is reasonable (not future, not too far past)
            if 2000 <= found_year <= current_year:
                tax_year = found_year
            else:
                # Default to previous year if extracted year is invalid
                tax_year = current_year - 1
        else:
            # No year found, default to previous year
            tax_year = current_year - 1
        
        # Extract key values using patterns
        parsed = ParsedTaxReturn(tax_year=tax_year)
        
        for field, pattern in self.TAX_FORM_PATTERNS['1040'].items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                if field == 'filing_status':
                    parsed.filing_status = match.group(1)
                else:
                    # Convert to float, removing commas
                    value = float(match.group(1).replace(',', ''))
                    setattr(parsed, field, value)
        
        # Calculate refund or owed
        if parsed.withholding > parsed.total_tax:
            parsed.refund_or_owed = parsed.withholding - parsed.total_tax
        else:
            parsed.refund_or_owed = -(parsed.total_tax - parsed.withholding)
        
        # Extract Schedule 1 income if present
        if 'schedule 1' in text.lower():
            self._parse_schedule_1(text, parsed)
        
        # Extract Schedule A deductions if present
        if 'schedule a' in text.lower():
            self._parse_schedule_a(text, parsed)
        
        return parsed
    
    async def parse_bank_statement(self, file_content: bytes, file_type: str = 'pdf') -> Tuple[ParsedAccount, List[ParsedTransaction]]:
        """
        Parse bank statement for account info and transactions
        """
        if file_type == 'pdf':
            text = self._extract_pdf_text(file_content)
            # Also try to extract tables
            transactions_df = self._extract_pdf_tables(file_content)
        else:
            text = ""
            transactions_df = None
        
        # Extract account information
        account = self._extract_account_info(text)
        
        # Extract transactions
        transactions = []
        
        if transactions_df is not None and not transactions_df.empty:
            # Parse from extracted table
            for _, row in transactions_df.iterrows():
                transaction = self._parse_transaction_row(row)
                if transaction:
                    transactions.append(transaction)
        else:
            # Fallback to regex parsing
            transactions = self._extract_transactions_regex(text)
        
        # Auto-categorize transactions
        for transaction in transactions:
            transaction.category = self._categorize_transaction(transaction.description)
        
        return account, transactions
    
    async def parse_investment_statement(self, file_content: bytes) -> Dict:
        """
        Parse investment/brokerage statement
        """
        text = self._extract_pdf_text(file_content)
        
        result = {
            'account_value': 0,
            'holdings': [],
            'transactions': [],
            'performance': {}
        }
        
        # Extract account value
        value_match = re.search(r'account value.*?\$?([\d,]+\.?\d*)', text, re.IGNORECASE)
        if value_match:
            result['account_value'] = float(value_match.group(1).replace(',', ''))
        
        # Extract holdings
        holdings_pattern = r'([A-Z]{1,5})\s+.*?([\d,]+)\s+shares?.*?\$?([\d,]+\.?\d*)'
        for match in re.finditer(holdings_pattern, text):
            result['holdings'].append({
                'symbol': match.group(1),
                'shares': float(match.group(2).replace(',', '')),
                'value': float(match.group(3).replace(',', ''))
            })
        
        # Extract YTD performance
        ytd_match = re.search(r'ytd.*?return.*?([\d\.]+)%', text, re.IGNORECASE)
        if ytd_match:
            result['performance']['ytd_return'] = float(ytd_match.group(1))
        
        return result
    
    async def parse_csv_transactions(self, file_content: bytes, format_type: str = 'auto') -> List[ParsedTransaction]:
        """
        Parse CSV transaction file (bank exports, Mint, etc.)
        """
        # Read CSV
        df = pd.read_csv(io.BytesIO(file_content))
        
        # Auto-detect format
        if format_type == 'auto':
            format_type = self._detect_csv_format(df)
        
        transactions = []
        
        # Map columns based on format
        column_mappings = {
            'chase': {
                'date': 'Transaction Date',
                'description': 'Description',
                'amount': 'Amount',
                'type': 'Type'
            },
            'bofa': {
                'date': 'Date',
                'description': 'Description',
                'amount': 'Amount',
                'balance': 'Running Bal.'
            },
            'mint': {
                'date': 'Date',
                'description': 'Description',
                'amount': 'Amount',
                'category': 'Category',
                'account': 'Account Name'
            },
            'generic': {
                'date': df.columns[0],  # Assume first column is date
                'description': df.columns[1],
                'amount': df.columns[2]
            }
        }
        
        mapping = column_mappings.get(format_type, column_mappings['generic'])
        
        for _, row in df.iterrows():
            try:
                # Parse date
                date_str = str(row[mapping['date']])
                trans_date = pd.to_datetime(date_str).to_pydatetime()
                
                # Parse amount
                amount_str = str(row[mapping['amount']])
                amount = float(amount_str.replace('$', '').replace(',', ''))
                
                # Determine transaction type
                trans_type = 'credit' if amount > 0 else 'debit'
                amount = abs(amount)
                
                # Create transaction
                transaction = ParsedTransaction(
                    date=trans_date,
                    description=str(row[mapping['description']]),
                    amount=amount,
                    transaction_type=trans_type,
                    category=row.get(mapping.get('category')) if 'category' in mapping else None,
                    account=row.get(mapping.get('account')) if 'account' in mapping else None
                )
                
                # Auto-categorize if not already categorized
                if not transaction.category:
                    transaction.category = self._categorize_transaction(transaction.description)
                
                # Extract merchant
                transaction.merchant = self._extract_merchant(transaction.description)
                
                transactions.append(transaction)
                
            except Exception as e:
                print(f"Error parsing row: {e}")
                continue
        
        return transactions
    
    async def parse_docx_financial_doc(self, file_content: bytes) -> Dict:
        """
        Parse DOCX financial documents (letters, summaries, etc.)
        """
        doc = Document(io.BytesIO(file_content))
        
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        tables_data = []
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        
        # Extract financial values
        values = {
            'amounts': [],
            'dates': [],
            'percentages': []
        }
        
        # Find dollar amounts
        amount_pattern = r'\$?([\d,]+\.?\d*)'
        for match in re.finditer(amount_pattern, text):
            try:
                values['amounts'].append(float(match.group(1).replace(',', '')))
            except (ValueError, IndexError) as e:
                # Skip invalid number formats
                logger.debug(f"Failed to parse amount '{match.group(1)}': {e}")
                continue
        
        # Find dates
        date_pattern = r'\d{1,2}/\d{1,2}/\d{2,4}'
        for match in re.finditer(date_pattern, text):
            values['dates'].append(match.group())
        
        # Find percentages
        percent_pattern = r'([\d\.]+)%'
        for match in re.finditer(percent_pattern, text):
            values['percentages'].append(float(match.group(1)))
        
        return {
            'text': text,
            'tables': tables_data,
            'extracted_values': values
        }
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    
    def _extract_pdf_tables(self, file_content: bytes) -> pd.DataFrame:
        """Extract tables from PDF using pdfplumber"""
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                all_tables = []
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            df = pd.DataFrame(table[1:], columns=table[0])
                            all_tables.append(df)
                
                if all_tables:
                    return pd.concat(all_tables, ignore_index=True)
        except (Exception) as e:
            # PDF table extraction failed, return empty DataFrame
            logger.warning(f"Failed to extract PDF tables: {e}")
            pass

        return pd.DataFrame()
    
    async def _extract_image_text(self, file_content: bytes, doc_type: str = None) -> str:
        """
        OCR text extraction from image using hybrid Tesseract + Claude Vision.

        Args:
            file_content: Image bytes
            doc_type: Document type hint (W2, 1099, bank_statement, etc.)

        Returns:
            Extracted text
        """
        from app.services.ocr_hybrid import get_hybrid_ocr

        hybrid_ocr = get_hybrid_ocr()

        try:
            text, engine, quality = await hybrid_ocr.extract_text(
                image_bytes=file_content,
                doc_type=doc_type
            )

            logger.info(
                f"OCR completed",
                extra={
                    "engine": engine,
                    "quality": quality,
                    "doc_type": doc_type,
                    "text_length": len(text)
                }
            )

            return text

        except Exception as e:
            logger.error(f"Hybrid OCR failed, falling back to basic Tesseract: {e}")
            # Final fallback to basic Tesseract
            image = Image.open(io.BytesIO(file_content))
            text = pytesseract.image_to_string(image)
            return text
    
    def _categorize_transaction(self, description: str) -> str:
        """Auto-categorize transaction based on description"""
        description_lower = description.lower()
        
        for category, keywords in self.CATEGORY_RULES.items():
            for keyword in keywords:
                if keyword in description_lower:
                    return category
        
        # Default categories based on amount patterns
        if 'transfer' in description_lower:
            return 'transfer'
        elif 'deposit' in description_lower:
            return 'income'
        elif 'payment' in description_lower:
            return 'payment'
        
        return 'other'
    
    def _extract_merchant(self, description: str) -> Optional[str]:
        """Extract merchant name from transaction description"""
        # Remove common transaction prefixes
        clean_desc = re.sub(r'^(purchase|payment|transfer|deposit|withdrawal)\s+', '', description, flags=re.IGNORECASE)
        
        # Extract first meaningful words (likely merchant)
        words = clean_desc.split()
        if words:
            # Take first 2-3 words as merchant
            merchant = ' '.join(words[:min(3, len(words))])
            return merchant.title()
        
        return None
    
    def _extract_account_info(self, text: str) -> ParsedAccount:
        """Extract account information from statement text"""
        account = ParsedAccount(
            account_name="",
            account_number="",
            account_type="checking",
            balance=0,
            institution="",
            as_of_date=datetime.now()
        )
        
        # Extract account number (masked)
        acct_match = re.search(r'account.*?(\*+\d{4})', text, re.IGNORECASE)
        if acct_match:
            account.account_number = acct_match.group(1)
        
        # Extract balance
        balance_match = re.search(r'(?:ending|current|available).*?balance.*?\$?([\d,]+\.?\d*)', text, re.IGNORECASE)
        if balance_match:
            account.balance = float(balance_match.group(1).replace(',', ''))
        
        # Extract institution (common bank names)
        banks = ['chase', 'bank of america', 'wells fargo', 'citibank', 'capital one', 'pnc', 'td bank']
        for bank in banks:
            if bank in text.lower():
                account.institution = bank.title()
                break
        
        # Extract statement date
        date_match = re.search(r'statement.*?date.*?(\d{1,2}/\d{1,2}/\d{2,4})', text, re.IGNORECASE)
        if date_match:
            account.as_of_date = datetime.strptime(date_match.group(1), '%m/%d/%Y')
        
        return account
    
    def _detect_csv_format(self, df: pd.DataFrame) -> str:
        """Auto-detect CSV format based on column names"""
        columns_lower = [col.lower() for col in df.columns]
        
        if 'transaction date' in columns_lower:
            return 'chase'
        elif 'running bal.' in columns_lower:
            return 'bofa'
        elif 'category' in columns_lower and 'account name' in columns_lower:
            return 'mint'
        else:
            return 'generic'
    
    def _parse_transaction_row(self, row: pd.Series) -> Optional[ParsedTransaction]:
        """Parse a single transaction row from dataframe"""
        try:
            # Find date column
            date_val = None
            for col in row.index:
                if 'date' in col.lower():
                    date_val = pd.to_datetime(row[col]).to_pydatetime()
                    break
            
            if not date_val:
                return None
            
            # Find amount column
            amount_val = None
            for col in row.index:
                if 'amount' in col.lower():
                    amount_str = str(row[col]).replace('$', '').replace(',', '')
                    amount_val = float(amount_str)
                    break
            
            if amount_val is None:
                return None
            
            # Find description
            desc_val = ""
            for col in row.index:
                if 'description' in col.lower() or 'memo' in col.lower():
                    desc_val = str(row[col])
                    break
            
            return ParsedTransaction(
                date=date_val,
                description=desc_val,
                amount=abs(amount_val),
                transaction_type='credit' if amount_val > 0 else 'debit'
            )
        except (ValueError, KeyError, TypeError) as e:
            # Failed to parse transaction row - skip invalid data
            logger.debug(f"Failed to parse transaction row: {e}")
            return None
    
    def _parse_schedule_1(self, text: str, parsed: ParsedTaxReturn):
        """Parse Schedule 1 additional income"""
        # Business income
        business_match = re.search(r'business income.*?\$?([\d,]+)', text, re.IGNORECASE)
        if business_match:
            parsed.business_income = float(business_match.group(1).replace(',', ''))
        
        # Rental income
        rental_match = re.search(r'rental.*?income.*?\$?([\d,]+)', text, re.IGNORECASE)
        if rental_match:
            parsed.rental_income = float(rental_match.group(1).replace(',', ''))
    
    def _parse_schedule_a(self, text: str, parsed: ParsedTaxReturn):
        """Parse Schedule A itemized deductions"""
        parsed.standard_or_itemized = "itemized"
        
        # Mortgage interest
        mortgage_match = re.search(r'home mortgage interest.*?\$?([\d,]+)', text, re.IGNORECASE)
        if mortgage_match:
            parsed.mortgage_interest = float(mortgage_match.group(1).replace(',', ''))
        
        # Charitable donations
        charity_match = re.search(r'gifts to charity.*?\$?([\d,]+)', text, re.IGNORECASE)
        if charity_match:
            parsed.charitable_donations = float(charity_match.group(1).replace(',', ''))
        
        # State and local taxes
        salt_match = re.search(r'state and local.*?taxes.*?\$?([\d,]+)', text, re.IGNORECASE)
        if salt_match:
            parsed.state_local_taxes = min(10000, float(salt_match.group(1).replace(',', '')))